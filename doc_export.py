"""Utilities to detect specially marked documentation blocks in assistant responses
and export them automatically to Word (.docx) files.

Markers:
    [WORD_DOC]
    (content)
    [/WORD_DOC]

Public functions:
- extraer_bloque_word(respuesta: str) -> str | None
- guardar_respuesta_en_word(texto_respuesta: str, ruta_salida: str) -> str
- procesar_respuesta(respuesta: str, ruta_salida: str | None = None) -> dict

Returns from procesar_respuesta:
    {
        'tiene_bloque': bool,
        'ruta_archivo': str | None,
        'contenido_extraido': str | None,
        'error': str | None
    }
"""
from __future__ import annotations
import os
import re
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple

try:
    import docx  # python-docx
except ImportError as e:  # pragma: no cover
    raise RuntimeError("La librería 'python-docx' no está instalada. Añádela a requirements.txt") from e

import logger

WORD_DOC_INICIO = r"\[WORD_DOC\]"
WORD_DOC_FIN = r"\[/WORD_DOC\]"
BLOQUE_PATRON = re.compile(WORD_DOC_INICIO + r"(.*?)" + WORD_DOC_FIN, re.DOTALL | re.IGNORECASE)


def extraer_bloque_word(respuesta: str) -> Optional[str]:
    """Extrae el contenido entre [WORD_DOC] y [/WORD_DOC].

    Args:
        respuesta: Texto completo devuelto por el asistente.

    Returns:
        Contenido interno (str) si se encuentra, o None si no existe el bloque.
    """
    if not respuesta:
        return None
    match = BLOQUE_PATRON.search(respuesta)
    if not match:
        return None
    contenido = match.group(1).strip()  # Limpiar espacios extremos
    logger.debug(f"Bloque [WORD_DOC] detectado. Longitud: {len(contenido)} caracteres", "doc_export.extraer_bloque_word")
    return contenido or None


MD_HEADER_RE = re.compile(r'^(#{1,6})\s+(.*)$')
MD_ULIST_RE = re.compile(r'^(?P<indent>\s*)[-*+]\s+(?P<item>.*)$')
MD_OLIST_RE = re.compile(r'^(?P<indent>\s*)\d+\.\s+(?P<item>.*)$')
MD_CODEBLOCK_FENCE_RE = re.compile(r'^```(.*)$')
MD_INLINE_CODE_RE = re.compile(r'`([^`]+)`')
MD_BOLD_RE = re.compile(r'(\*\*|__)(.+?)\1')
MD_ITALIC_RE = re.compile(r'(?<!\*)\*(?!\*)([^*]+)(?<!\*)\*(?!\*)|_(?!_)([^_]+)_(?!_)')


def _ensure_code_style(document: 'docx.document.Document') -> None:
    """Crea un estilo monoespaciado simple para código si no existe."""
    from docx.shared import Pt
    styles = document.styles
    if 'Code' in [s.name for s in styles]:
        return
    style = styles.add_style('Code', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10)


TABLE_SEPARATOR_RE = re.compile(r'^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$')
TABLE_ROW_RE = re.compile(r'^\s*\|(.+)\|\s*$')
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
BLOCKQUOTE_RE = re.compile(r'^>\s?(.*)$')


def _normalize_links(text: str) -> str:
    """Convierte [texto](url) a 'texto (url)' para preservarlo en Word si no hacemos estilo especial."""
    def repl(m):
        label, url = m.group(1).strip(), m.group(2).strip()
        return f"{label} ({url})"
    return LINK_RE.sub(repl, text)


def _add_table(document: 'docx.document.Document', header: List[str], rows: List[List[str]]) -> None:
    table = document.add_table(rows=len(rows)+1, cols=len(header))
    table.style = 'Light List Accent 1' if 'Light List Accent 1' in [s.name for s in document.styles] else table.style
    # Header
    for j, h in enumerate(header):
        table.cell(0, j).text = h.strip()
    # Rows
    for i, row in enumerate(rows, start=1):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell.strip()


def _parse_table_lines(buffer: List[str]) -> Optional[Tuple[List[str], List[List[str]]]]:
    if len(buffer) < 2:
        return None
    # First line header, second separator
    header_match = TABLE_ROW_RE.match(buffer[0])
    sep_ok = bool(TABLE_SEPARATOR_RE.match(buffer[1]))
    if not header_match or not sep_ok:
        return None
    header = [c.strip() for c in header_match.group(1).split('|')]
    body: List[List[str]] = []
    for line in buffer[2:]:
        m = TABLE_ROW_RE.match(line)
        if not m:
            return None
        body.append([c.strip() for c in m.group(1).split('|')])
    return header, body


def _add_markdown_paragraph(document: 'docx.document.Document', line: str, list_ctx: dict) -> None:
    """Añade una línea interpretando Markdown extendido (listas, tablas, citas, etc)."""
    stripped = line.strip()

    # Encabezados
    m = MD_HEADER_RE.match(line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        document.add_heading(text, level=min(level, 6))
        return

    # Code fence start/end
    if MD_CODEBLOCK_FENCE_RE.match(line):
        if not list_ctx.get('in_code_block'):
            list_ctx['in_code_block'] = True
            list_ctx['code_buffer'] = []
        else:
            # flush code block
            para = document.add_paragraph('\n'.join(list_ctx['code_buffer']))
            para.style = 'Code'
            list_ctx['in_code_block'] = False
            list_ctx['code_buffer'] = []
        return

    if list_ctx.get('in_code_block'):
        list_ctx['code_buffer'].append(line)
        return

    # Eliminar separadores horizontales o líneas solo con signos repetidos
    if stripped and len(stripped) >= 3 and set(stripped).issubset({'-', '_', '*'}) and len(set(stripped)) == 1:
        return

    # Citas
    m_quote = BLOCKQUOTE_RE.match(line)
    if m_quote:
        quote_text = m_quote.group(1)
        para = document.add_paragraph()
        run = para.add_run(f"❝ {quote_text}")
        run.italic = True
        return

    # Tablas: detect if starting a table block
    if line.strip().startswith('|') and line.strip().endswith('|'):
        # Acumular líneas hasta que deje de ser tabla
        tbl_lines = [line]
        list_ctx['pending_table'] = True
        list_ctx['table_buffer'] = tbl_lines
        return
    if list_ctx.get('pending_table'):
        if line.strip().startswith('|') and line.strip().endswith('|') or TABLE_SEPARATOR_RE.match(line):
            list_ctx['table_buffer'].append(line)
            return
        # Intentar parsear
        parsed = _parse_table_lines(list_ctx['table_buffer'])
        if parsed:
            header, body = parsed
            _add_table(document, header, body)
        else:
            # Si falla, volcar como texto plano
            for l in list_ctx['table_buffer']:
                document.add_paragraph(_normalize_links(l))
        list_ctx['pending_table'] = False
        list_ctx['table_buffer'] = []
        # Continuar procesando la línea actual como normal (no parte de tabla)

    # Omitir líneas vacías para respetar el espaciado del estilo Word
    if not stripped:
        return

    # Listas con indentación
    m_ul = MD_ULIST_RE.match(line)
    m_ol = MD_OLIST_RE.match(line)
    if m_ul or m_ol:
        indent = (m_ul or m_ol).group('indent')
        item = (m_ul or m_ol).group('item').strip()
        bullet = '•' if m_ul else '-'  # prefijo simple
        # Calcular nivel (4 espacios ~ 1 nivel)
        level = len(indent.replace('\t', '    ')) // 2  # un poco más compacto
        prefix = '   ' * level + bullet + ' '
        para = document.add_paragraph()
        _add_inline_runs(para, prefix + item)
        return

    # Línea en blanco
    if not line.strip():
        document.add_paragraph('')
        return

    # Párrafo normal con inline formatting
    para = document.add_paragraph()
    _add_inline_runs(para, _normalize_links(line))


def _apply_inline_formatting(run: 'docx.text.run.Run') -> None:
    text = run.text
    # Bold / Italic / Inline code no separables sin recrear runs; se maneja en _add_inline_runs
    pass


def _add_inline_runs(para: 'docx.text.paragraph.Paragraph', line: str) -> None:
    """Fragmenta la línea creando runs para bold, italic e inline code."""
    # Primero reemplazamos inline code conservando delimitadores para procesar en orden
    tokens: List[Tuple[str, str]] = []  # (tipo, texto)

    # Procesar inline code dividiendo por backticks
    def split_inline_code(s: str) -> List[Tuple[str, str]]:
        parts = re.split(r'(`[^`]+`)', s)
        out: List[Tuple[str, str]] = []
        for p in parts:
            if not p:
                continue
            if p.startswith('`') and p.endswith('`'):
                out.append(('code', p[1:-1]))
            else:
                out.append(('text', p))
        return out

    code_split = split_inline_code(line)

    def split_bold_italic(segment: str) -> List[Tuple[str, str]]:
        # Proceso iterativo sobre bold y luego italic
        result: List[Tuple[str, str]] = []
        pos = 0
        pattern = re.compile(r'(\*\*.+?\*\*|__.+?__|\*(?!\*)([^*]+)\*(?!\*)|_(?!_)([^_]+)_(?!_))')
        for m in pattern.finditer(segment):
            if m.start() > pos:
                result.append(('text', segment[pos:m.start()]))
            token = m.group(0)
            if token.startswith('**') or token.startswith('__'):
                inner = token[2:-2]
                result.append(('bold', inner))
            elif token.startswith('*') or token.startswith('_'):
                inner = token[1:-1]
                result.append(('italic', inner))
            pos = m.end()
        if pos < len(segment):
            result.append(('text', segment[pos:]))
        return result

    for kind, content in code_split:
        if kind == 'code':
            run = para.add_run(content)
            run.style = 'Code'
            continue
        # Split bold/italic inside normal text
        for kind2, seg in split_bold_italic(content):
            run = para.add_run(seg)
            if kind2 == 'bold':
                run.bold = True
            elif kind2 == 'italic':
                run.italic = True


def guardar_respuesta_en_word(texto_respuesta: str, ruta_salida: str) -> str:
    """Genera un archivo .docx aplicando un subconjunto de formato Markdown.

        Soporta:
            - Encabezados (# .. ######)
            - Listas (anidadas por indentación) no ordenadas (- * +) y ordenadas (n.)
            - Citas (> texto)
            - Tablas Markdown sencillas (| col1 | col2 |)
            - Enlaces [texto](url) -> "texto (url)"
            - Código en bloque (``` triple backticks)
            - Código inline (`code`)
            - Bold (** ** o __ __) e italic (* * o _ _)
    """
    if not texto_respuesta:
        raise ValueError("texto_respuesta vacío")

    os.makedirs(ruta_salida, exist_ok=True)

    ahora = datetime.now()
    nombre_archivo = f"documentacion_{ahora.strftime('%Y%m%d_%H%M%S')}.docx"
    ruta_archivo = os.path.abspath(os.path.join(ruta_salida, nombre_archivo))

    doc = docx.Document()
    _ensure_code_style(doc)
    doc.add_heading('Documentación del asistente', level=1)
    doc.add_paragraph(f"Generado: {ahora.strftime('%Y-%m-%d %H:%M:%S')}")

    list_ctx: Dict[str, Any] = {"in_code_block": False, "code_buffer": []}
    for linea in texto_respuesta.splitlines():
        _add_markdown_paragraph(doc, linea.rstrip('\r'), list_ctx)

    # Si quedó un bloque de código abierto, volcarlo
    if list_ctx.get('in_code_block') and list_ctx.get('code_buffer'):
        para = doc.add_paragraph('\n'.join(list_ctx['code_buffer']))
        para.style = 'Code'

    try:
        doc.save(ruta_archivo)
    except Exception as e:  # pragma: no cover
        logger.error(f"Error guardando archivo Word: {e}", "doc_export.guardar_respuesta_en_word")
        raise

    logger.info(f"Archivo Word generado: {ruta_archivo}", "doc_export.guardar_respuesta_en_word")
    return ruta_archivo


def procesar_respuesta(respuesta: str, ruta_salida: Optional[str] = None) -> Dict[str, Any]:
    """Detecta si la respuesta contiene bloque Word y lo exporta.

    Args:
        respuesta: Texto completo del asistente.
        ruta_salida: Directorio destino (por defecto 'data/word_docs').

    Returns:
        Diccionario con metadatos del proceso.
    """
    ruta_salida = ruta_salida or os.environ.get('WORD_DOC_OUTPUT_DIR', os.path.join('data', 'word_docs'))

    resultado = {
        'tiene_bloque': False,
        'ruta_archivo': None,
        'contenido_extraido': None,
        'error': None
    }

    try:
        contenido = extraer_bloque_word(respuesta)
        if not contenido:
            logger.debug("No se encontró bloque [WORD_DOC] en la respuesta", "doc_export.procesar_respuesta")
            return resultado

        resultado['tiene_bloque'] = True
        resultado['contenido_extraido'] = contenido
        ruta_archivo = guardar_respuesta_en_word(contenido, ruta_salida)
        resultado['ruta_archivo'] = ruta_archivo
        return resultado
    except Exception as e:  # pragma: no cover
        resultado['error'] = str(e)
        logger.error(f"Fallo procesando respuesta para Word: {e}", "doc_export.procesar_respuesta")
        return resultado


# Ejemplo de uso manual
if __name__ == '__main__':  # pragma: no cover
    ejemplo = """Texto cualquiera antes\n[WORD_DOC]\n# Título Ejemplo\nContenido de prueba\nCon varias líneas\n\nMás texto.\n[/WORD_DOC]\nTexto luego"""
    info = procesar_respuesta(ejemplo)
    print(info)
